# src/processors/specialized_parsers.py
"""
Specialized File Parsers for Nexus Cognition System
Implements robust parsers for each supported file type with advanced error handling
"""

import pandas as pd
import numpy as np
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from abc import ABC, abstractmethod
import re
import io
import json
import logging
from datetime import datetime
import hashlib

# Import processor-specific libraries
try:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_LIBS_AVAILABLE = True
except ImportError:
    PDF_LIBS_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .bulletproof_pipeline import DocumentChunk, ChunkType, ProcessingMethod

logger = logging.getLogger(__name__)

class BaseProcessor(ABC):
    """Base class for all file processors"""
    
    def __init__(self):
        self.chunk_id_counter = 0
    
    @abstractmethod
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse file and return list of document chunks"""
        pass
    
    def _generate_chunk_id(self, prefix: str = "") -> str:
        """Generate unique chunk ID"""
        self.chunk_id_counter += 1
        base = f"{prefix}_{self.chunk_id_counter}" if prefix else str(self.chunk_id_counter)
        return hashlib.md5(base.encode()).hexdigest()[:8]
    
    def _create_chunk(
        self, 
        content: str, 
        chunk_type: ChunkType, 
        source_location: str,
        confidence: float = 1.0,
        metadata: Optional[Dict] = None
    ) -> DocumentChunk:
        """Helper to create document chunks"""
        return DocumentChunk(
            chunk_id=self._generate_chunk_id(),
            type=chunk_type,
            content=content.strip(),
            source_location=source_location,
            confidence=confidence,
            metadata=metadata or {},
            processing_method=ProcessingMethod.NATIVE
        )

class ATFProcessor(BaseProcessor):
    """
    Advanced ATF (Axon Text Format) Parser
    Handles scientific data format with proper header/data separation
    """
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse ATF file with robust error handling"""
        try:
            return self._parse_atf_file(file_path)
        except Exception as e:
            logger.error(f"ATF parsing failed: {e}")
            # Fallback to basic text parsing
            return self._fallback_text_parse(file_path)
    
    def _parse_atf_file(self, file_path: str) -> List[DocumentChunk]:
        """Main ATF parsing logic"""
        chunks = []
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        
        if not lines:
            raise ValueError("Empty ATF file")
        
        # Validate ATF format
        if not lines[0].strip().startswith("ATF"):
            raise ValueError("Invalid ATF format - missing ATF header")
        
        # Parse ATF structure
        atf_info = self._parse_atf_header(lines)
        header_metadata = self._extract_header_metadata(lines, atf_info)
        data_df = self._extract_data_section(lines, atf_info)
        
        # Create metadata chunk
        metadata_content = self._format_metadata(header_metadata, atf_info)
        chunks.append(self._create_chunk(
            content=metadata_content,
            chunk_type=ChunkType.METADATA,
            source_location="ATF_header",
            metadata={
                'atf_version': atf_info.get('version'),
                'header_lines': atf_info.get('header_lines'),
                'total_metadata_fields': len(header_metadata)
            }
        ))
        
        # Create data chunks
        if not data_df.empty:
            data_chunks = self._create_data_chunks(data_df, atf_info)
            chunks.extend(data_chunks)
        
        # Create summary chunk
        summary = self._create_atf_summary(header_metadata, data_df, atf_info)
        chunks.append(self._create_chunk(
            content=summary,
            chunk_type=ChunkType.TEXT,
            source_location="ATF_summary",
            metadata={'chunk_role': 'summary'}
        ))
        
        logger.info(f"ATF parsing successful: {len(chunks)} chunks created")
        return chunks
    
    def _parse_atf_header(self, lines: List[str]) -> Dict[str, Any]:
        """Parse ATF header structure"""
        atf_line = lines[0].strip()
        
        # Extract version (e.g., "ATF 1.0")
        version_match = re.search(r'ATF\s+([\d.]+)', atf_line)
        version = version_match.group(1) if version_match else "unknown"
        
        # Find header size
        header_lines = 0
        data_starts_at = 2  # Default fallback
        
        if len(lines) > 1:
            header_line = lines[1].strip()
            header_match = re.search(r'(\d+)\s+header\s+lines?', header_line, re.IGNORECASE)
            
            if header_match:
                header_lines = int(header_match.group(1))
                data_starts_at = header_lines + 2  # +2 for ATF line and header count line
            else:
                # Try to find data section automatically
                for i, line in enumerate(lines[2:], start=2):
                    if line.strip() and not '=' in line and '\t' in line:
                        # Likely start of data section
                        data_starts_at = i
                        header_lines = i - 2
                        break
        
        return {
            'version': version,
            'header_lines': header_lines,
            'data_starts_at': data_starts_at,
            'total_lines': len(lines)
        }
    
    def _extract_header_metadata(self, lines: List[str], atf_info: Dict) -> Dict[str, str]:
        """Extract metadata from ATF header section"""
        metadata = {}
        
        # Parse metadata lines (between line 2 and data start)
        for i in range(2, min(atf_info['data_starts_at'], len(lines))):
            line = lines[i].strip().replace('"', '')
            
            if '=' in line:
                try:
                    key, value = line.split('=', 1)
                    metadata[key.strip()] = value.strip()
                except ValueError:
                    # Handle malformed metadata lines
                    metadata[f'raw_line_{i}'] = line
            elif line:  # Non-empty line without '='
                metadata[f'note_line_{i}'] = line
        
        return metadata
    
    def _extract_data_section(self, lines: List[str], atf_info: Dict) -> pd.DataFrame:
        """Extract and parse the data section"""
        data_start = atf_info['data_starts_at']
        
        if data_start >= len(lines):
            logger.warning("No data section found in ATF file")
            return pd.DataFrame()
        
        # Find column headers (usually the line just before data)
        column_line_idx = data_start - 1
        if column_line_idx < len(lines):
            column_names = lines[column_line_idx].strip().replace('"', '').split('\t')
        else:
            column_names = None
        
        # Extract data lines
        data_lines = lines[data_start:]
        data_content = '\n'.join(data_lines)
        
        if not data_content.strip():
            return pd.DataFrame()
        
        try:
            # Parse data with pandas
            df = pd.read_csv(
                io.StringIO(data_content),
                delimiter='\t',
                names=column_names,
                header=None,
                na_values=['', 'NA', 'N/A', 'null', 'NULL'],
                keep_default_na=True,
                encoding='utf-8'
            )
            
            # Clean up the dataframe
            df = df.dropna(how='all')  # Remove completely empty rows
            
            # Convert numeric columns
            for col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='ignore')
            
            logger.info(f"ATF data extracted: {len(df)} rows, {len(df.columns)} columns")
            return df
            
        except Exception as e:
            logger.error(f"Failed to parse ATF data section: {e}")
            return pd.DataFrame()
    
    def _create_data_chunks(self, df: pd.DataFrame, atf_info: Dict) -> List[DocumentChunk]:
        """Create contextual chunks from ATF data"""
        chunks = []
        
        # Strategy: Create meaningful row-based chunks with context
        for idx, row in df.iterrows():
            # Convert row to readable text
            row_items = []
            for col, value in row.items():
                if pd.notna(value):
                    row_items.append(f"{col}: {value}")
            
            if row_items:
                content = f"ATF Data Record {idx + 1}: " + ", ".join(row_items)
                
                chunks.append(self._create_chunk(
                    content=content,
                    chunk_type=ChunkType.TABLE,
                    source_location=f"data_row_{idx + 1}",
                    metadata={
                        'row_index': idx,
                        'column_count': len(row_items),
                        'data_type': 'atf_record'
                    }
                ))
        
        # Create column summary chunks
        for col in df.columns:
            if df[col].dtype in ['int64', 'float64']:
                stats = df[col].describe()
                content = f"ATF Column {col} Statistics: Mean={stats['mean']:.2f}, Std={stats['std']:.2f}, Min={stats['min']}, Max={stats['max']}"
                
                chunks.append(self._create_chunk(
                    content=content,
                    chunk_type=ChunkType.METADATA,
                    source_location=f"column_stats_{col}",
                    metadata={
                        'column_name': col,
                        'data_type': str(df[col].dtype),
                        'statistics': stats.to_dict()
                    }
                ))
        
        return chunks
    
    def _format_metadata(self, metadata: Dict[str, str], atf_info: Dict) -> str:
        """Format metadata into readable text"""
        lines = [f"ATF File Version: {atf_info['version']}"]
        lines.append(f"Header Lines: {atf_info['header_lines']}")
        
        for key, value in metadata.items():
            lines.append(f"{key}: {value}")
        
        return "\n".join(lines)
    
    def _create_atf_summary(self, metadata: Dict, data_df: pd.DataFrame, atf_info: Dict) -> str:
        """Create comprehensive summary of ATF file"""
        summary_parts = [
            f"ATF File Summary (Version {atf_info['version']})",
            f"Metadata fields: {len(metadata)}",
            f"Data records: {len(data_df) if not data_df.empty else 0}",
            f"Data columns: {len(data_df.columns) if not data_df.empty else 0}"
        ]
        
        if not data_df.empty:
            summary_parts.extend([
                f"Numeric columns: {len(data_df.select_dtypes(include=[np.number]).columns)}",
                f"Text columns: {len(data_df.select_dtypes(include=['object']).columns)}"
            ])
        
        # Add key metadata
        important_keys = ['title', 'description', 'date', 'author', 'experiment']
        for key in important_keys:
            if key in metadata:
                summary_parts.append(f"{key.title()}: {metadata[key]}")
        
        return "\n".join(summary_parts)
    
    def _fallback_text_parse(self, file_path: str) -> List[DocumentChunk]:
        """Fallback to basic text parsing if ATF parsing fails"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return [self._create_chunk(
            content=content,
            chunk_type=ChunkType.TEXT,
            source_location="full_file",
            confidence=0.5,  # Lower confidence for fallback
            metadata={'parsing_method': 'fallback_text'}
        )]

class ExcelProcessor(BaseProcessor):
    """Advanced Excel processor with smart chunking"""
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse Excel file with multiple strategies"""
        chunks = []
        
        try:
            # Try openpyxl first for detailed analysis
            if OPENPYXL_AVAILABLE:
                chunks.extend(self._parse_with_openpyxl(file_path))
            else:
                # Fallback to pandas
                chunks.extend(self._parse_with_pandas(file_path))
        except Exception as e:
            logger.error(f"Excel parsing failed: {e}")
            # Last resort: try to read as CSV
            try:
                df = pd.read_excel(file_path, engine='openpyxl')
                chunks.extend(self._dataframe_to_chunks(df, "Sheet1"))
            except Exception as e2:
                logger.error(f"Excel fallback parsing also failed: {e2}")
                raise e
        
        return chunks
    
    def _parse_with_openpyxl(self, file_path: str) -> List[DocumentChunk]:
        """Parse using openpyxl for rich metadata"""
        chunks = []
        wb = load_workbook(file_path, read_only=True, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            # Convert to DataFrame for easier processing
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            
            if data:
                df = pd.DataFrame(data)
                # Clean up - first row might be headers
                if len(df) > 1:
                    df.columns = df.iloc[0]
                    df = df.drop(df.index[0]).reset_index(drop=True)
                
                chunks.extend(self._dataframe_to_chunks(df, sheet_name))
        
        wb.close()
        return chunks
    
    def _parse_with_pandas(self, file_path: str) -> List[DocumentChunk]:
        """Parse using pandas for speed"""
        chunks = []
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            chunks.extend(self._dataframe_to_chunks(df, sheet_name))
        
        return chunks
    
    def _dataframe_to_chunks(self, df: pd.DataFrame, sheet_name: str) -> List[DocumentChunk]:
        """Convert DataFrame to contextual chunks"""
        chunks = []
        
        if df.empty:
            return chunks
        
        # Create header chunk
        if not df.columns.empty:
            header_content = f"Excel Sheet '{sheet_name}' contains columns: " + ", ".join(str(col) for col in df.columns)
            chunks.append(self._create_chunk(
                content=header_content,
                chunk_type=ChunkType.HEADER,
                source_location=f"{sheet_name}!headers",
                metadata={'sheet_name': sheet_name, 'column_count': len(df.columns)}
            ))
        
        # Create row chunks (contextual format)
        for idx, row in df.iterrows():
            row_items = []
            for col, value in row.items():
                if pd.notna(value) and str(value).strip():
                    row_items.append(f"{col}: {value}")
            
            if row_items:
                content = f"Row {idx + 1} from {sheet_name}: " + ", ".join(row_items)
                
                chunks.append(self._create_chunk(
                    content=content,
                    chunk_type=ChunkType.TABLE,
                    source_location=f"{sheet_name}!R{idx + 1}",
                    metadata={
                        'sheet_name': sheet_name,
                        'row_index': idx,
                        'cell_count': len(row_items)
                    }
                ))
        
        # Create summary chunk for the sheet
        summary = f"Excel Sheet '{sheet_name}' Summary: {len(df)} rows, {len(df.columns)} columns"
        if len(df.select_dtypes(include=[np.number]).columns) > 0:
            summary += f", {len(df.select_dtypes(include=[np.number]).columns)} numeric columns"
        
        chunks.append(self._create_chunk(
            content=summary,
            chunk_type=ChunkType.METADATA,
            source_location=f"{sheet_name}!summary",
            metadata={'sheet_name': sheet_name, 'chunk_role': 'summary'}
        ))
        
        return chunks

class CSVProcessor(BaseProcessor):
    """Robust CSV processor with encoding detection"""
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse CSV with multiple delimiter detection"""
        
        # Detect delimiter and encoding
        delimiter, encoding = self._detect_csv_properties(file_path)
        
        try:
            df = pd.read_csv(
                file_path, 
                delimiter=delimiter, 
                encoding=encoding,
                na_values=['', 'NA', 'N/A', 'null', 'NULL'],
                keep_default_na=True
            )
            
            return self._dataframe_to_chunks(df, "CSV")
            
        except Exception as e:
            logger.error(f"CSV parsing failed: {e}")
            # Fallback to simple text parsing
            return self._fallback_text_parse(file_path)
    
    def _detect_csv_properties(self, file_path: str) -> tuple[str, str]:
        """Detect CSV delimiter and encoding"""
        import csv
        
        # Detect encoding
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] if encoding_result['confidence'] > 0.8 else 'utf-8'
        except ImportError:
            encoding = 'utf-8'
        
        # Detect delimiter
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
        except Exception:
            delimiter = ','  # Default fallback
        
        return delimiter, encoding
    
    def _dataframe_to_chunks(self, df: pd.DataFrame, source_name: str) -> List[DocumentChunk]:
        """Convert CSV DataFrame to chunks"""
        chunks = []
        
        # Similar to Excel processing but adapted for CSV
        if not df.empty:
            # Header chunk
            header_content = f"CSV file contains columns: " + ", ".join(str(col) for col in df.columns)
            chunks.append(self._create_chunk(
                content=header_content,
                chunk_type=ChunkType.HEADER,
                source_location="csv_headers",
                metadata={'column_count': len(df.columns)}
            ))
            
            # Row chunks
            for idx, row in df.iterrows():
                row_items = []
                for col, value in row.items():
                    if pd.notna(value) and str(value).strip():
                        row_items.append(f"{col}: {value}")
                
                if row_items:
                    content = f"CSV Row {idx + 1}: " + ", ".join(row_items)
                    
                    chunks.append(self._create_chunk(
                        content=content,
                        chunk_type=ChunkType.TABLE,
                        source_location=f"row_{idx + 1}",
                        metadata={'row_index': idx, 'cell_count': len(row_items)}
                    ))
        
        return chunks
    
    def _fallback_text_parse(self, file_path: str) -> List[DocumentChunk]:
        """Fallback text parsing for problematic CSVs"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        
        chunks = []
        for i, line in enumerate(lines[:100]):  # Limit to first 100 lines
            if line.strip():
                chunks.append(self._create_chunk(
                    content=line.strip(),
                    chunk_type=ChunkType.TEXT,
                    source_location=f"line_{i + 1}",
                    confidence=0.7,
                    metadata={'parsing_method': 'fallback_text'}
                ))
        
        return chunks

class PDFProcessor(BaseProcessor):
    """Advanced PDF processor with text extraction"""
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse PDF with multiple extraction methods"""
        
        if not PDF_LIBS_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        chunks = []
        
        try:
            # Try PyMuPDF first (fitz)
            chunks = self._parse_with_fitz(file_path)
        except Exception as e:
            logger.warning(f"PyMuPDF parsing failed: {e}")
            try:
                # Fallback to PyPDF2
                chunks = self._parse_with_pypdf2(file_path)
            except Exception as e2:
                logger.error(f"PyPDF2 parsing also failed: {e2}")
                raise e2
        
        return chunks
    
    def _parse_with_fitz(self, file_path: str) -> List[DocumentChunk]:
        """Parse using PyMuPDF (fitz)"""
        chunks = []
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            
            if text.strip():
                # Split into paragraphs
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                
                for para_idx, paragraph in enumerate(paragraphs):
                    if len(paragraph) > 20:  # Only meaningful chunks
                        chunks.append(self._create_chunk(
                            content=paragraph,
                            chunk_type=ChunkType.TEXT,
                            source_location=f"page_{page_num + 1}_para_{para_idx + 1}",
                            metadata={
                                'page_number': page_num + 1,
                                'paragraph_index': para_idx + 1,
                                'extraction_method': 'fitz'
                            }
                        ))
        
        doc.close()
        return chunks
    
    def _parse_with_pypdf2(self, file_path: str) -> List[DocumentChunk]:
        """Parse using PyPDF2"""
        chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    chunks.append(self._create_chunk(
                        content=text.strip(),
                        chunk_type=ChunkType.TEXT,
                        source_location=f"page_{page_num + 1}",
                        metadata={
                            'page_number': page_num + 1,
                            'extraction_method': 'pypdf2'
                        }
                    ))
        
        return chunks

class DOCXProcessor(BaseProcessor):
    """DOCX document processor"""
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse DOCX file"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        chunks = []
        doc = Document(file_path)
        
        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                chunks.append(self._create_chunk(
                    content=paragraph.text.strip(),
                    chunk_type=ChunkType.TEXT,
                    source_location=f"paragraph_{para_idx + 1}",
                    metadata={
                        'paragraph_index': para_idx + 1,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    }
                ))
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                table_text = "\n".join(["\t".join(row) for row in table_data])
                chunks.append(self._create_chunk(
                    content=f"Table {table_idx + 1}:\n{table_text}",
                    chunk_type=ChunkType.TABLE,
                    source_location=f"table_{table_idx + 1}",
                    metadata={
                        'table_index': table_idx + 1,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    }
                ))
        
        return chunks

class TextProcessor(BaseProcessor):
    """Simple text file processor"""
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """Parse plain text file"""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        chunks = []
        
        for idx, paragraph in enumerate(paragraphs):
            if len(paragraph) > 10:  # Only meaningful chunks
                chunks.append(self._create_chunk(
                    content=paragraph,
                    chunk_type=ChunkType.TEXT,
                    source_location=f"paragraph_{idx + 1}",
                    metadata={'paragraph_index': idx + 1}
                ))
        
        return chunks


# Export all processors
__all__ = [
    'BaseProcessor', 'ATFProcessor', 'ExcelProcessor', 'CSVProcessor',
    'PDFProcessor', 'DOCXProcessor', 'TextProcessor'
]